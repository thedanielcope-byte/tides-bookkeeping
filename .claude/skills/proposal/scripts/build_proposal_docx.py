#!/usr/bin/env python3
"""
Tides Bookkeeping — branded engagement-proposal .docx generator.

Usage:
    python3 build_proposal_docx.py content.json "/path/out.docx" [/path/to/logo.png]

content.json schema (all keys optional except client_short; sensible defaults applied):
{
  "client_short": "Pizza Bar",
  "eyebrow": "ENGAGEMENT PROPOSAL",
  "eyebrow_sub": "Books Assessment & Cleanup",
  "prepared_for": ["Jeffery Kohn — Owner", "Pizza Bar LLC", "1627 Collins Ave, Miami Beach FL", "Platform: Restaurant365 · Paycor"],
  "prepared_by":  ["Dan Cope", "Tides Bookkeeping", "dan@tidesbookkeeping.com", "August 11, 2026"],
  "exec_summary": ["para 1", "para 2"],
  "headline_box": {"lead": "The headline: ", "body": "...", "close": "..."},
  "reviewed": "Working directly inside your Restaurant365 environment, we examined ...",
  "findings": [["Title", "Description."], ...],
  "why_it_matters": [["Lead phrase ", "rest of sentence."], ...],
  "phases": [
     {"tag":"PHASE 1","title":"Catch-Up & Cleanup (one-time)","price":"$7,500",
      "pricenote":"billed 50% to start, 50% on completion","items":["...","..."],"note":"optional footnote"},
     {"tag":"PHASE 2","title":"Ongoing Monthly Bookkeeping","price":"$1,495 / month",
      "pricenote":"month-to-month, no long-term contract","items":["...","..."]}
  ],
  "investment_rows": [["Phase 1 — Cleanup","One-time (50/50)","$7,500"], ["First-year total","...","$25,440"]],
  "investment_note": "For comparison, current spend already runs ~$X/month for ...",
  "payment_schedule": {"intro":"The two fees cover different periods ...",
      "rows":[["On signing","$3,750","Cleanup deposit (50%) — Jan–Aug"], ...], "note":"No long-term contract ..."},
  "cpa": {"title":"Working With Your CPA","body":"..."},
  "why_tides": [["Restaurant-specialized. ","..."], ...],
  "next_steps": ["step 1", "step 2", "step 3"],
  "cta": "Ready to get your books right?",
  "footer": "Tides Bookkeeping · Confidential — prepared for <client> · Pricing valid 30 days",
  "email": "dan@tidesbookkeeping.com",
  "site": "tidesbookkeeping.com"
}
"""
import sys, json, os, tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0E,0x3A,0x53); TEAL=RGBColor(0x1F,0x8A,0x8A); LIGHT=RGBColor(0x5B,0x6B,0x73)
INK=RGBColor(0x22,0x2A,0x2E); WHITE=RGBColor(0xFF,0xFF,0xFF)

def trimmed_logo(src):
    """Trim whitespace around the Tides logo, composite on white; return temp path or None."""
    try:
        from PIL import Image
        im=Image.open(src).convert("RGBA"); px=im.load(); w,h=im.size
        mask=Image.new("L",(w,h),0); mp=mask.load()
        for y in range(h):
            for x in range(w):
                r,g,b,a=px[x,y]
                if a>10 and not (r>245 and g>245 and b>245): mp[x,y]=255
        bb=mask.getbbox()
        if not bb: return None
        crop=im.crop(bb); pad=24
        canvas=Image.new("RGBA",(crop.width+pad*2,crop.height+pad*2),(255,255,255,255))
        canvas.alpha_composite(crop,(pad,pad))
        out=os.path.join(tempfile.gettempdir(),"tides_logo_docx.png")
        canvas.convert("RGB").save(out); return out
    except Exception:
        return None

def main():
    content_path, out_path = sys.argv[1], sys.argv[2]
    logo_src = sys.argv[3] if len(sys.argv)>3 else "logo.png"
    C=json.load(open(content_path))
    LOGO = trimmed_logo(logo_src) if os.path.exists(logo_src) else None

    doc=Document()
    base=doc.styles['Normal']; base.font.name='Calibri'; base.font.size=Pt(10.5)
    base.font.color.rgb=INK; base.paragraph_format.space_after=Pt(6); base.paragraph_format.line_spacing=1.12

    def shade(cell,hexc):
        tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
        sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
    def cmargins(cell,t=60,b=60,l=100,r=100):
        tcPr=cell._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
        for tag,val in (('top',t),('bottom',b),('start',l),('end',r)):
            e=OxmlElement(f'w:{tag}'); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
        tcPr.append(m)
    def no_border(t):
        bd=OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'none'); bd.append(e)
        t._tbl.tblPr.append(bd)
    def para(text='',size=10.5,color=INK,bold=False,italic=False,after=6,before=0,align=None):
        p=doc.add_paragraph()
        if align: p.alignment=align
        p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
        if text:
            r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color
        return p
    def heading(text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        r=p.add_run(text); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=NAVY
        pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
        bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6'); bt.set(qn('w:space'),'4'); bt.set(qn('w:color'),'1F8A8A')
        pbdr.append(bt); pPr.append(pbdr)
    def bullet(lead,text):
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
        if lead:
            r=p.add_run(lead); r.font.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=NAVY
        r2=p.add_run(text); r2.font.size=Pt(10.5)

    # HEADER
    h=doc.add_table(rows=1,cols=2); no_border(h); c=h.rows[0].cells
    cmargins(c[0],l=20,t=40,b=20); cmargins(c[1],r=40,t=90)
    if LOGO: c[0].paragraphs[0].add_run().add_picture(LOGO,width=Inches(2.35))
    else:
        rp=c[0].paragraphs[0].add_run('TIDES BOOKKEEPING'); rp.font.bold=True; rp.font.size=Pt(18); rp.font.color.rgb=NAVY
    tag=c[0].add_paragraph(); tag.paragraph_format.space_before=Pt(2)
    rt=tag.add_run('Restaurant & Hospitality Accounting'); rt.font.size=Pt(8.5); rt.font.color.rgb=TEAL; rt.font.bold=True
    p2=c[1].paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p2.add_run(C.get('eyebrow','ENGAGEMENT PROPOSAL')); r.font.bold=True; r.font.size=Pt(12.5); r.font.color.rgb=NAVY
    p3=c[1].add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p3.add_run(C.get('eyebrow_sub','Books Assessment & Cleanup')); r.font.size=Pt(10); r.font.color.rgb=TEAL
    rule=doc.add_paragraph(); rule.paragraph_format.space_before=Pt(4); rule.paragraph_format.space_after=Pt(2)
    pPr=rule._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'18'); bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),'0E3A53')
    pbdr.append(bt); pPr.append(pbdr)

    # PREPARED FOR / BY
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    meta=doc.add_table(rows=1,cols=2); no_border(meta); mc=meta.rows[0].cells
    cmargins(mc[0],l=40); cmargins(mc[1],l=40)
    def kv(cell,label,lines):
        p=cell.paragraphs[0]; r=p.add_run(label); r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=TEAL
        p.paragraph_format.space_after=Pt(1)
        for i,ln in enumerate(lines):
            pp=cell.add_paragraph(); rr=pp.add_run(ln); rr.font.size=Pt(10.5)
            rr.font.color.rgb=NAVY if i==0 else LIGHT; rr.font.bold=(i==0); pp.paragraph_format.space_after=Pt(0)
    kv(mc[0],'PREPARED FOR',C.get('prepared_for',['Client']))
    kv(mc[1],'PREPARED BY',C.get('prepared_by',['Dan Cope','Tides Bookkeeping','dan@tidesbookkeeping.com']))

    # EXEC SUMMARY
    if C.get('exec_summary'):
        heading('Executive Summary')
        for pr in C['exec_summary']: para(pr)
    if C.get('headline_box'):
        hb=C['headline_box']; box=doc.add_table(rows=1,cols=1); no_border(box); bc=box.rows[0].cells[0]
        shade(bc,'F1F7F7'); cmargins(bc,t=140,b=140,l=180,r=180)
        bp=bc.paragraphs[0]
        if hb.get('lead'):
            r=bp.add_run(hb['lead']); r.font.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=NAVY
        r2=bp.add_run(hb.get('body','')); r2.font.size=Pt(10.5); r2.font.color.rgb=INK
        if hb.get('close'):
            bp2=bc.add_paragraph(); bp2.paragraph_format.space_before=Pt(6)
            r=bp2.add_run(hb['close']); r.font.size=Pt(10.5); r.font.bold=True; r.font.color.rgb=NAVY

    if C.get('reviewed'):
        heading('What We Reviewed'); para(C['reviewed'])

    # FINDINGS
    if C.get('findings'):
        heading('Current-State Findings')
        for i,(t,d) in enumerate(C['findings'],1):
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
            r=p.add_run(f'{i}.  '); r.font.bold=True; r.font.color.rgb=TEAL; r.font.size=Pt(10.5)
            r2=p.add_run(t+'.  '); r2.font.bold=True; r2.font.color.rgb=NAVY; r2.font.size=Pt(10.5)
            r3=p.add_run(d); r3.font.size=Pt(10.5)

    if C.get('why_it_matters'):
        heading('Why It Matters')
        for lead,text in C['why_it_matters']: bullet(lead,text)

    # PHASES
    if C.get('phases'):
        heading('Proposed Engagement')
        for ph in C['phases']:
            t=doc.add_table(rows=1,cols=1); no_border(t); cell=t.rows[0].cells[0]
            cmargins(cell,t=120,b=120,l=160,r=160)
            tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
            for edge,w,col in (('top','24','1F8A8A'),('left','8','E3ECEC'),('right','8','E3ECEC'),('bottom','8','E3ECEC')):
                e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),w); e.set(qn('w:space'),'0'); e.set(qn('w:color'),col); b.append(e)
            tcPr.append(b)
            p=cell.paragraphs[0]
            r=p.add_run(ph.get('tag','')+'   '); r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=TEAL
            r2=p.add_run(ph.get('title','')); r2.font.bold=True; r2.font.size=Pt(12.5); r2.font.color.rgb=NAVY
            pp=cell.add_paragraph(); pp.paragraph_format.space_after=Pt(6)
            rp=pp.add_run(ph.get('price','')); rp.font.bold=True; rp.font.size=Pt(15); rp.font.color.rgb=NAVY
            if ph.get('pricenote'):
                rn=pp.add_run('   '+ph['pricenote']); rn.font.size=Pt(9.5); rn.font.color.rgb=LIGHT; rn.font.italic=True
            for it in ph.get('items',[]):
                b2=cell.add_paragraph(style='List Bullet'); b2.paragraph_format.space_after=Pt(2); b2.paragraph_format.left_indent=Inches(0.28)
                rb=b2.add_run(it); rb.font.size=Pt(10)
            if ph.get('note'):
                np=cell.add_paragraph(); np.paragraph_format.space_before=Pt(4)
                rn=np.add_run(ph['note']); rn.font.size=Pt(9); rn.font.italic=True; rn.font.color.rgb=LIGHT
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

    def money_table(headers, rows, highlight_last=False):
        tbl=doc.add_table(rows=len(rows)+1,cols=len(headers)); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
        hd=tbl.rows[0].cells
        for i,l in enumerate(headers):
            shade(hd[i],'0E3A53'); cmargins(hd[i])
            p=hd[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(l); r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for ri,row in enumerate(rows,1):
            cells=tbl.rows[ri].cells; fill='F1F7F7' if ri%2 else 'FFFFFF'
            if highlight_last and ri==len(rows): fill='E8F0EF'
            for j,val in enumerate(row):
                shade(cells[j],fill); cmargins(cells[j])
                p=cells[j].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                r=p.add_run(val); r.font.size=Pt(9.5 if len(headers)==3 and j==2 and len(val)>20 else 10)
                r.font.bold=(highlight_last and ri==len(rows)) or (j==len(headers)-1)
                r.font.color.rgb=NAVY if ((highlight_last and ri==len(rows)) or j==len(headers)-1) else INK

    if C.get('investment_rows'):
        heading('Investment Summary')
        money_table(['Service','Structure','Investment'], C['investment_rows'], highlight_last=True)
        if C.get('investment_note'):
            para(C['investment_note'], size=9.5, italic=True, color=LIGHT, before=6)

    if C.get('payment_schedule'):
        ps=C['payment_schedule']; heading('Payment Schedule')
        if ps.get('intro'): para(ps['intro'], after=8)
        money_table(['When','Amount','What it covers'], ps.get('rows',[]))
        if ps.get('note'): para(ps['note'], size=9.5, italic=True, color=LIGHT, before=6)

    if C.get('cpa'):
        heading(C['cpa'].get('title','Working With Your CPA')); para(C['cpa'].get('body',''))

    if C.get('why_tides'):
        heading('Why Tides Bookkeeping')
        for lead,text in C['why_tides']: bullet(lead,text)

    if C.get('next_steps'):
        heading('Next Steps')
        for i,s in enumerate(C['next_steps'],1):
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
            r=p.add_run(f'{i}.  '); r.font.bold=True; r.font.color.rgb=TEAL; r.font.size=Pt(10.5)
            r2=p.add_run(s); r2.font.size=Pt(10.5)

    if C.get('cta'):
        para(' ',after=4)
        cta=doc.add_table(rows=1,cols=1); no_border(cta); cc=cta.rows[0].cells[0]
        shade(cc,'0E3A53'); cmargins(cc,t=140,b=140,l=180,r=180)
        p=cc.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(C['cta']+'  '); r.font.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=WHITE
        contact=f"{C.get('email','dan@tidesbookkeeping.com')}  ·  {C.get('site','tidesbookkeeping.com')}"
        r2=p.add_run(contact); r2.font.size=Pt(10.5); r2.font.color.rgb=RGBColor(0xC9,0xE3,0xE3)

    foot=doc.sections[0].footer; fp=foot.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run(C.get('footer','Tides Bookkeeping · Confidential')); r.font.size=Pt(8); r.font.color.rgb=LIGHT

    doc.save(out_path); print('Saved:', out_path)

if __name__=='__main__':
    main()
